from __future__ import annotations

import re
from pathlib import Path
from typing import TYPE_CHECKING

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

if TYPE_CHECKING:
    from app.settings import Settings


REPORT_DOCX_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def generate_report_docx(report: dict, settings: "Settings") -> dict:
    report_id = safe_report_id(str(report.get("report_id") or "report"))
    output_dir = reports_dir(settings)
    output_dir.mkdir(parents=True, exist_ok=True)
    path = output_dir / f"{report_id}.docx"

    document = Document()
    configure_document(document)

    context = report.get("structured_context") or {}
    add_title(document, "单警工作日报")
    add_meta_table(
        document,
        [
            ("日报ID", report_id),
            ("任务编号", report.get("mission_id") or ""),
            ("报告类型", report.get("report_type") or "daily"),
            ("生成时间", report.get("generated_at") or ""),
            ("生成模型", report.get("model") or ""),
            ("生成后端", report.get("backend") or ""),
            ("复核要求", "需人工复核后归档" if report.get("requires_human_confirmation", True) else "无需复核"),
        ],
    )

    add_heading(document, "基础信息", level=1)
    add_meta_table(
        document,
        [
            ("日期", report_date_text(report.get("generated_at"))),
            ("值班民警", report.get("officer_name") or ""),
            ("警号", report.get("operator_id") or ""),
            ("设备/小脑", report.get("device_id") or settings.device_id),
            ("班次", ""),
            ("巡逻区域", ""),
        ],
    )

    add_heading(document, "今日工作情况", level=1)
    summary = context.get("multisource_summary") or {}
    counts = summary.get("media_counts") or {}
    add_bullet(document, f"接警数量：")
    add_bullet(document, f"处警数量：")
    add_bullet(document, f"巡逻时长：")
    add_bullet(document, f"重点情况：{summary.get('overall') or format_counts(counts)}")

    add_heading(document, "多源材料汇总", level=1)
    add_multisource_table(document, summary)

    add_heading(document, "警情记录", level=1)
    add_paragraph(document, f"案件编号：{case_id_for_report(report.get('generated_at'))}", bold=True)
    add_meta_table(
        document,
        [
            ("警情类型", ""),
            ("时间", report.get("generated_at") or ""),
            ("地点", ""),
            ("涉及人员", ""),
        ],
    )
    add_heading(document, "现场情况", level=2)
    add_plain_lines(document, normalize_report_text(report.get("content") or ""))
    add_heading(document, "处置结果", level=2)
    add_plain_lines(
        document,
        "本日报由边缘小脑根据视频、录音和图片材料生成，为 AI 草稿。现场情况、处置经过及最终结果需由值班民警结合原始附件复核后确认。",
    )

    add_heading(document, "附件信息", level=1)
    add_attachment_table(document, context.get("media_items") or [])

    add_heading(document, "备注", level=1)
    add_plain_lines(
        document,
        "AI 识别结果仅作为候选提示，不得直接作为最终事实认定。附件 SHA-256 已写入结构化上下文，可用于后续证据链核验。",
    )

    add_heading(document, "签字", level=1)
    add_meta_table(document, [("值班民警", ""), ("审核人", "")])

    document.save(path)
    return {
        "format": "docx",
        "file_name": path.name,
        "file_path": str(path),
        "download_url": f"/api/v1/reports/{report_id}/download",
        "content_type": REPORT_DOCX_CONTENT_TYPE,
    }


def reports_dir(settings: "Settings") -> Path:
    return settings.data_dir / "reports"


def safe_report_id(value: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9_.-]+", "-", value).strip(".-")
    return cleaned[:120] or "report"


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "SimSun"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(4)

    for style_name, size, color in [
        ("Title", 18, RGBColor(31, 78, 121)),
        ("Heading 1", 13, RGBColor(31, 78, 121)),
        ("Heading 2", 11.5, RGBColor(47, 84, 150)),
    ]:
        style = styles[style_name]
        style.font.name = "SimHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.color.rgb = color


def add_title(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="Title")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run(text)


def add_heading(document: Document, text: str, level: int) -> None:
    document.add_heading(text, level=level)


def add_paragraph(document: Document, text: str, bold: bool = False) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = bold


def add_bullet(document: Document, text: str) -> None:
    document.add_paragraph(text, style="List Bullet")


def add_plain_lines(document: Document, text: str) -> None:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    if not lines:
        document.add_paragraph("")
        return
    for line in lines:
        document.add_paragraph(strip_markdown_line(line))


def add_meta_table(document: Document, rows: list[tuple[str, object]]) -> None:
    table = document.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = str(label)
        cells[1].text = str(value or "")
        shade_cell(cells[0], "EAF2F8")
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_table_widths(table, [1800, 7000])
    document.add_paragraph()


def add_multisource_table(document: Document, summary: dict) -> None:
    rows = [
        ("视频", join_items(summary.get("video_summaries") or [])),
        ("录音", join_items(summary.get("audio_summaries") or [])),
        ("图片", join_items(summary.get("image_summaries") or [])),
        ("综合判断", summary.get("overall") or "未形成可汇总的多源材料结论。"),
    ]
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = table.rows[0].cells
    headers[0].text = "信息源"
    headers[1].text = "摘要"
    for cell in headers:
        shade_cell(cell, "D9EAF7")
    for source, text in rows:
        cells = table.add_row().cells
        cells[0].text = source
        cells[1].text = text or "未纳入该类材料。"
        cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_table_widths(table, [1500, 7300])
    document.add_paragraph()


def add_attachment_table(document: Document, media_items: list[dict]) -> None:
    table = document.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = table.rows[0].cells
    for cell, title in zip(headers, ["类型", "文件名", "证据ID", "分析摘要"], strict=True):
        cell.text = title
        shade_cell(cell, "D9EAF7")
    for item in media_items:
        cells = table.add_row().cells
        cells[0].text = media_type_label(item.get("media_type") or "")
        cells[1].text = str(item.get("source_name") or item.get("source_uri") or "")
        cells[2].text = str(item.get("evidence_id") or "")
        cells[3].text = str(item.get("summary") or media_summary(item))[:600]
    if not media_items:
        cells = table.add_row().cells
        cells[0].text = "-"
        cells[1].text = "未纳入附件"
    set_table_widths(table, [1200, 2800, 1800, 3000])
    document.add_paragraph()


def set_table_widths(table, widths: list[int]) -> None:
    for row in table.rows:
        for index, width in enumerate(widths):
            if index < len(row.cells):
                row.cells[index].width = Inches(width / 1440)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def normalize_report_text(text: str) -> str:
    return "\n".join(strip_markdown_line(line) for line in text.splitlines())


def strip_markdown_line(line: str) -> str:
    cleaned = line.strip()
    cleaned = re.sub(r"^#{1,6}\s*", "", cleaned)
    cleaned = re.sub(r"^\|?\s*-{2,}.*$", "", cleaned)
    cleaned = re.sub(r"^\|\s*", "", cleaned)
    cleaned = re.sub(r"\s*\|\s*", "：", cleaned).strip("： ")
    cleaned = re.sub(r"^[-*]\s+", "", cleaned)
    return cleaned


def join_items(items: list[object]) -> str:
    values = [str(item).strip() for item in items if str(item).strip()]
    return "\n".join(values[:8])


def media_summary(item: dict) -> str:
    analysis = item.get("analysis") or {}
    return str(analysis.get("structured_text") or analysis.get("transcript") or analysis.get("status") or "")


def media_type_label(media_type: str) -> str:
    return {
        "video": "视频",
        "audio": "录音",
        "image": "图片",
        "other": "其他",
        "unknown": "未知",
    }.get(media_type, media_type)


def format_counts(counts: dict) -> str:
    if not counts:
        return "未纳入视频、录音或图片材料。"
    return "；".join(f"{media_type_label(key)} {value} 个" for key, value in counts.items())


def report_date_text(value: object) -> str:
    text = str(value or "")
    return text[:10] if len(text) >= 10 else ""


def case_id_for_report(value: object) -> str:
    date_text = report_date_text(value).replace("-", "")
    return f"{date_text or '00000000'}-001"
